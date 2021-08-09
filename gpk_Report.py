from gpk_utilities import *
from gpk_Score import *

import copy
import docx
import os
import matplotlib.pyplot as plt
import pandas as pd
from dateutil import tz


class gpk_Report():
    def __init__(self,Profile,doc_name,path = None, timezone = "Asia/Shanghai"):
        self.Profile = Profile
        self.tz = tz.gettz(timezone)
        self.document  = docx.Document()
        self.doc_name = doc_name
        if path is None:
            self.path = os.getcwd() + "/Report_Pics"
        #Create a Directory 
        try:
            os.makedirs(self.path)
            print("Directory " , self.path ,  " Created ")
        except FileExistsError:
            print("Directory " , self.path ,  " already exists")  
        
    def Pics_Gen(self):
        global D
        #Generate Pictures to be added to the plot 
        Analysis = DF_Analysis(self.Profile.todos.Archive)  
    ##Prepare Data:
        #-WEEKDAYS List of avaliable wkdays since last Monday
        date_LM = str(Last_monday())
        #df = Analysis.SEARCH('date_done',lambda date: DATE(date) >= DATE(date_LM))
        DF = Analysis.df
        DF = DF.sort_values('date_done', key = lambda L: [DATE(x) for x in L],ascending = False)
        df = pd.DataFrame((DF.groupby(by = 'date_done').Reward.agg(sum)))
        
        dates = list(df.index)
        rewards = list(df['Reward'])
        Dict = {date:reward for date,reward in zip(dates,rewards)}
        Dict = Fill_date(Dict)
                
        WEEKDAYS = list([D[DATE(date).weekday()+1] for date in Dict.keys()])
        #-PRODUCTIVITY List of of avaliable Rewards since last Monday
        PRODUCTIVITY = list(Dict.values())
        #WEEK_PRODUCTIVITY/WEEK_NUM
        def Cal_Week(Date_0,date):
                    delta = DATE(Date_0) - DATE(date)
                    return -round((delta.days/7))
        Most_recent = DF.iloc[0]['date_done']
        Last_n_df = copy.copy(DF)
        Last_n_df['Week'] = [Cal_Week(Most_recent, date) for date in list(Last_n_df['date_done'])] 
        
        temp = Last_n_df.groupby('Week').Reward.agg(sum)
        res = pd.DataFrame(temp).sort_values(by = 'Week', key = lambda L: [int(i) for i in L],ascending = False)
        #Finally:
        WEEK_NUM = [i for i in res.index]
        WEEK_PRODUCTIVITY = res['Reward']
            
        
        
##Ploting:
    #fig_all.png
        fig_all = plt.figure(figsize=(20,10))
        plt.bar(range(len(PRODUCTIVITY)),PRODUCTIVITY)
        fig_all.suptitle('Productivity of ALL time')
        plt.xlabel('Dates')
        plt.ylabel('Productivity/credits received')
        plt.savefig(self.path+'/fig_all.png')

    #fig_7.png
        fig_last_7 = plt.figure(figsize=(20,10))
        plt.bar(WEEKDAYS[-7:len(WEEKDAYS)],PRODUCTIVITY[-7:len(WEEKDAYS)])
        fig_last_7.suptitle('Productivity in Last 7 days')
        plt.xlabel('Dates')
        plt.ylabel('Productivity/credits received')
        plt.savefig(self.path+'/fig_7.png')
    #fig_wk.png
        fig_weeks = plt.figure(figsize=(20,10))
        plt.bar(WEEK_NUM,WEEK_PRODUCTIVITY)
        fig_weeks.suptitle('Productivity by weeks')
        plt.xlabel('Weeks')
        plt.ylabel('Productivity/credits received')
        plt.savefig(self.path+'/fig_wk.png')
    #okr_pie1.png
        okr_pie1 = plt.figure(figsize=(15,10))
        okr_pie1.suptitle('Time by Sections')
        Analysis.Plot_Sec(plt_method = plt.pie,kargs = 
                          {'x': 'X_new','explode':'explode',
                           'labels' : 'Labs', 
                          'autopct' : 'lambda value: str(round(value,2))+"%"'})
        plt.savefig(self.path+'/okr_pie1.png')
    #okr_pie2.png
        okr_pie2 = plt.figure(figsize=(15,10))
        okr_pie2.suptitle('Productivity by Sections')
        Analysis.Plot_Sec(sec ='Reward',plt_method = plt.pie,kargs = 
                          {'x': 'X_new','explode':'explode',
                           'labels' : 'Labs', 
                          'autopct' : 'lambda value: str(round(value,2))+"%"'})
        plt.savefig(self.path+'/okr_pie2.png')
    
        #plt.show()
        
    def df_to_docx_table(self,df, document = None,mutate = True):
        "Mudate the current docx by adding a Table"
        if document is None:
            document = self.document
        if not mutate:
            document = copy.copy(document)
        # add table ------------------
        table = document.add_table(1, cols=df.shape[1])
        # populate header row --------
        heading_cells = table.rows[0].cells
        for idx in range(df.shape[1]):
            heading_cells[idx].text = df.columns[idx]
        # add a data row for each item
        for row in df.index:
            cells = table.add_row().cells
            for j,col in enumerate(df.columns):
                text = df.at[row,col]
                cells[j].text = str(text)
        if not mutate:
            return document 

        
    def OKR_report(self,season):   
        #Create Document 
        doc = self.document
        time_stamp = str(datetime.datetime.now(self.tz))
    ##Menu##:
        doc.add_heading('OKR Report:{}'.format(season), 0)
        doc.add_heading(time_stamp, 2)
        doc.add_heading("Menu:", 1)
        paraMenu = doc.add_paragraph("Part1: OKR Progress\nPart2:OKRLOG\nPart3:OKR Stats\nPart4:Inventory History\nPart5:Evaluation",style='Intense Quote')
        doc.add_heading("Part1:OKR Progress", 1)
        
    #Part 1: OKR Week Progress 
        paraOKR_Progress = doc.add_paragraph(print_collector(
            self.Profile.todos.Load.WeekObjective.show))
        doc.add_page_break()
        
    #Part 2: OKR Log 
        doc.add_heading("Part2:OKR Log", 1)
        #Add A Data Frame
        df = self.Profile.todos.Archive 
        Analysis = DF_Analysis(df)
        df = Analysis.Last_n_day(n=7)
        df = df.drop(['ObjectID','KeyResult ID','Task Category',
                      'date_done','week_day','description','Deadline'],axis = 1)
        self.df_to_docx_table(df)
        doc.add_page_break()
        
    #Part 3: OKR Statistics 
        self.Pics_Gen()
        
        doc.add_heading("Part3:OKR Statistics", 1)
        doc.add_picture(self.path+'/fig_7.png', width=docx.shared.Cm(15),height=docx.shared.Cm(7))
        doc.add_picture(self.path+'/fig_wk.png', width=docx.shared.Cm(15),height=docx.shared.Cm(7))
        doc.add_picture(self.path+'/fig_all.png', width=docx.shared.Cm(15),height=docx.shared.Cm(7))
        doc.add_picture(self.path+'/okr_pie1.png', width=docx.shared.Cm(11),height=docx.shared.Cm(7))
        doc.add_picture(self.path+'/okr_pie2.png', width=docx.shared.Cm(11),height=docx.shared.Cm(7))
        doc.add_page_break()
        
    # #Part 4: OKR Inventory History
    #     doc.add_heading("Part4:OKR Inventory History", 1)
    #     paraOKR_bag_hist = doc.add_paragraph(get_inventory_history(data,datetime.datetime.now(China_tz).date()))
    #     doc.add_page_break()
    
    # Part 5:OKR Weekly Evaluation
        Analysis = DF_Analysis(self.Profile.todos.Archive,(14,10))
        Analysis.Rest_fig((14,10))
        Analysis.Plot_Score(self.Profile.todos.Load_backup)
        fig = Analysis.get_fig()
        fig.savefig(self.path + "/fig_score.png")
        
        doc.add_heading("Part5:OKR Weekly Evaluation", 1)
        score = score_okr(self.Profile.todos.Load) 
        grade = grade_okr(score)
        summary = "Total Score:{}\nLetter Grade:{}".format(score,grade)
        paraSummary = doc.add_paragraph(summary,style='Intense Quote')
        doc.add_picture(self.path+'/fig_score.png', width=docx.shared.Cm(14),height=docx.shared.Cm(10))
        
    #Finally:
        not_saved = True
        while not_saved:
            try:
                self.save_path = tk.filedialog.asksaveasfilename(initialdir = self.path ,
                                                                 initialfile = self.doc_name,
                                 title = "Where would you like the report to be saved?" ,
                                 filetypes = (("word document","*.docx"),("all files","*.*")))
                doc.save(self.save_path + '.docx')
                print('Report Generated')
                not_saved = False
            except:
                self.doc_name += '(1)'
                
    
        
if __name__ == '__main__':
    #MrFAKE_user_file.gpk
    User_name = 'Leo_TEST'##
    file_path = f"D:\GPK\gpk_saves\\{User_name}_user_file.gpk"
    with open(file_path,'rb') as INfile:
        Profile = pickle.load(INfile)
    Report = gpk_Report(Profile,'OKR_Report_S3_W2')
    Report.OKR_report('Season 3 Week2')